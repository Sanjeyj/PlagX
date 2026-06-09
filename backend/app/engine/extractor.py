"""
Text Extraction Engine
Extracts text from PDF, DOCX, and TXT files while preserving
paragraph structure, character offsets, and page information.
"""

import io
import re
import logging
from pathlib import Path
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class SentenceInfo:
    """A single sentence with position metadata."""
    sentence_index: int
    text: str
    start_char: int
    end_char: int


@dataclass
class ParagraphInfo:
    """A single paragraph with position metadata."""
    text: str
    paragraph_index: int
    page_number: int
    start_char: int
    end_char: int
    word_count: int
    sentences: list[SentenceInfo] = field(default_factory=list)


@dataclass
class ExtractionResult:
    """Complete extraction result with structure metadata."""
    full_text: str
    paragraphs: list[ParagraphInfo] = field(default_factory=list)
    total_words: int = 0
    total_pages: int = 1
    total_paragraphs: int = 0
    source_filename: str = ""
    file_type: str = ""


class TextExtractor:
    """
    Multi-format text extraction engine.
    Supports PDF (native + OCR), DOCX, and TXT.
    Preserves paragraph structure and character offsets.
    """

    SUPPORTED_TYPES = {".pdf", ".docx", ".txt"}

    def extract(self, file_path: str | Path) -> ExtractionResult:
        """Extract text from a file based on its extension."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info(f"Extracting text from {path.name} (type: {ext})")

        if ext == ".pdf":
            result = self._extract_pdf(path)
        elif ext == ".docx":
            result = self._extract_docx(path)
        elif ext == ".txt":
            result = self._extract_txt(path)
        else:
            raise ValueError(f"Unsupported: {ext}")

        result.source_filename = path.name
        result.file_type = ext
        result.total_words = len(result.full_text.split())
        result.total_paragraphs = len(result.paragraphs)

        logger.info(
            f"Extracted {result.total_words} words, "
            f"{result.total_paragraphs} paragraphs, "
            f"{result.total_pages} pages"
        )
        return result

    def _extract_pdf(self, path: Path) -> ExtractionResult:
        """Extract text from PDF using PyMuPDF with pdfplumber fallback."""
        try:
            return self._extract_pdf_pymupdf(path)
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed, trying pdfplumber: {e}")
            try:
                return self._extract_pdf_pdfplumber(path)
            except Exception as e2:
                logger.warning(f"pdfplumber failed, trying OCR: {e2}")
                return self._extract_pdf_ocr(path)

    def _extract_pdf_pymupdf(self, path: Path) -> ExtractionResult:
        """Primary PDF extraction using PyMuPDF (fitz)."""
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        paragraphs = []
        full_text_parts = []
        char_offset = 0

        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks", sort=True)

            for block in blocks:
                if block[6] == 0:  # Text block (not image)
                    text = block[4].strip()
                    if not text:
                        continue

                    # Split block into paragraphs
                    for para_text in text.split("\n\n"):
                        para_text = para_text.strip()
                        if not para_text:
                            continue

                        # Clean up excessive whitespace within paragraph
                        para_text = re.sub(r'\s+', ' ', para_text)

                        para = ParagraphInfo(
                            text=para_text,
                            paragraph_index=len(paragraphs),
                            page_number=page_num + 1,
                            start_char=char_offset,
                            end_char=char_offset + len(para_text),
                            word_count=len(para_text.split()),
                        )
                        paragraphs.append(para)
                        full_text_parts.append(para_text)
                        char_offset += len(para_text) + 1  # +1 for newline separator

        doc.close()
        full_text = "\n".join(full_text_parts)

        return ExtractionResult(
            full_text=full_text,
            paragraphs=paragraphs,
            total_pages=len(doc) if hasattr(doc, '__len__') else doc.page_count,
        )

    def _extract_pdf_pdfplumber(self, path: Path) -> ExtractionResult:
        """Fallback PDF extraction using pdfplumber."""
        import pdfplumber

        paragraphs = []
        full_text_parts = []
        char_offset = 0

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                for para_text in text.split("\n\n"):
                    para_text = re.sub(r'\s+', ' ', para_text.strip())
                    if not para_text:
                        continue

                    para = ParagraphInfo(
                        text=para_text,
                        paragraph_index=len(paragraphs),
                        page_number=page_num + 1,
                        start_char=char_offset,
                        end_char=char_offset + len(para_text),
                        word_count=len(para_text.split()),
                    )
                    paragraphs.append(para)
                    full_text_parts.append(para_text)
                    char_offset += len(para_text) + 1

            total_pages = len(pdf.pages)

        return ExtractionResult(
            full_text="\n".join(full_text_parts),
            paragraphs=paragraphs,
            total_pages=total_pages,
        )

    def _extract_pdf_ocr(self, path: Path) -> ExtractionResult:
        """OCR fallback for scanned PDFs using Tesseract."""
        try:
            import fitz
            import pytesseract
            from PIL import Image

            doc = fitz.open(str(path))
            paragraphs = []
            full_text_parts = []
            char_offset = 0

            for page_num in range(len(doc)):
                page = doc[page_num]
                # Render page to image
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                # OCR the image
                text = pytesseract.image_to_string(img)

                for para_text in text.split("\n\n"):
                    para_text = re.sub(r'\s+', ' ', para_text.strip())
                    if not para_text or len(para_text) < 10:
                        continue

                    para = ParagraphInfo(
                        text=para_text,
                        paragraph_index=len(paragraphs),
                        page_number=page_num + 1,
                        start_char=char_offset,
                        end_char=char_offset + len(para_text),
                        word_count=len(para_text.split()),
                    )
                    paragraphs.append(para)
                    full_text_parts.append(para_text)
                    char_offset += len(para_text) + 1

            doc.close()
            return ExtractionResult(
                full_text="\n".join(full_text_parts),
                paragraphs=paragraphs,
                total_pages=len(doc),
            )
        except ImportError:
            logger.error("Tesseract OCR not available")
            raise RuntimeError("OCR dependencies not installed (pytesseract, Pillow)")

    def _extract_docx(self, path: Path) -> ExtractionResult:
        """Extract text from DOCX preserving paragraph structure, tables, and sentences."""
        import tempfile
        import shutil
        import zipfile
        import unicodedata
        import spacy
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        logger.info("DOCX upload started")

        # Step 1: Safe File Handling (Copy to tempfile)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            with open(path, "rb") as src:
                shutil.copyfileobj(src, tmp)
            temp_path = tmp.name

        # Step 2: ZIP validation
        if not zipfile.is_zipfile(temp_path):
            import os
            os.remove(temp_path)
            raise ValueError("Invalid or corrupted DOCX file. Please upload a valid Microsoft Word document.")
        logger.info("ZIP validation passed")

        # Load SpaCy for sentence parsing
        try:
            nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            logger.warning(f"Failed to load spaCy model: {e}")
            nlp = None

        text_blocks = []

        # Step 3 & 4: Extraction with Fallback
        try:
            doc = Document(temp_path)
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_blocks.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                text_blocks.append(para.text)
            logger.info("python-docx extraction succeeded")
        except (ValueError, PackageNotFoundError, Exception) as e:
            logger.warning(f"python-docx failed ({e}), attempting mammoth fallback")
            try:
                import mammoth
                with open(temp_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    raw_text = result.value
                    if raw_text:
                        for block in raw_text.split("\n\n"):
                            if block.strip():
                                text_blocks.append(block)
                logger.info("mammoth extraction succeeded")
            except Exception as mammoth_e:
                import os
                os.remove(temp_path)
                logger.error(f"Both docx and mammoth extraction failed: {mammoth_e}")
                raise ValueError("Invalid or corrupted DOCX file. Please upload a valid Microsoft Word document.")
        
        # Cleanup temp file
        import os
        try:
            os.remove(temp_path)
        except Exception:
            pass

        paragraphs = []
        full_text_parts = []
        char_offset = 0
        page_number = 1
        words_on_page = 0
        sentence_count = 0

        for text in text_blocks:
            text = text.strip()
            if not text:
                continue

            # Step 8: Text sanitization
            text = unicodedata.normalize("NFKC", text)
            text = re.sub(r'\s+', ' ', text)
            word_count = len(text.split())

            # Estimate page breaks (~300 words per page)
            words_on_page += word_count
            if words_on_page > 300:
                page_number += 1
                words_on_page = word_count

            sentences_info = []
            if nlp:
                parsed_doc = nlp(text)
                for sent in parsed_doc.sents:
                    sent_text = sent.text.strip()
                    if sent_text:
                        local_start = text.find(sent_text)
                        if local_start == -1:
                            local_start = 0
                        
                        s_info = SentenceInfo(
                            sentence_index=sentence_count,
                            text=sent_text,
                            start_char=char_offset + local_start,
                            end_char=char_offset + local_start + len(sent_text)
                        )
                        sentences_info.append(s_info)
                        sentence_count += 1

            para_info = ParagraphInfo(
                text=text,
                paragraph_index=len(paragraphs),
                page_number=page_number,
                start_char=char_offset,
                end_char=char_offset + len(text),
                word_count=word_count,
                sentences=sentences_info
            )
            paragraphs.append(para_info)
            full_text_parts.append(text)
            char_offset += len(text) + 1  # +1 for newline separator

        logger.info(f"Extracted {len(paragraphs)} paragraphs")
        logger.info(f"Extracted {sentence_count} sentences")

        return ExtractionResult(
            full_text="\n".join(full_text_parts),
            paragraphs=paragraphs,
            total_pages=page_number,
        )

    def _extract_txt(self, path: Path) -> ExtractionResult:
        """Extract text from TXT with encoding detection."""
        import chardet

        raw = path.read_bytes()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8") or "utf-8"

        text = raw.decode(encoding, errors="replace")
        paragraphs = []
        full_text_parts = []
        char_offset = 0
        page_number = 1
        words_on_page = 0

        for para_text in re.split(r'\n\s*\n', text):
            para_text = re.sub(r'\s+', ' ', para_text.strip())
            if not para_text:
                continue

            word_count = len(para_text.split())
            words_on_page += word_count
            if words_on_page > 300:
                page_number += 1
                words_on_page = word_count

            para = ParagraphInfo(
                text=para_text,
                paragraph_index=len(paragraphs),
                page_number=page_number,
                start_char=char_offset,
                end_char=char_offset + len(para_text),
                word_count=word_count,
            )
            paragraphs.append(para)
            full_text_parts.append(para_text)
            char_offset += len(para_text) + 1

        return ExtractionResult(
            full_text="\n".join(full_text_parts),
            paragraphs=paragraphs,
            total_pages=page_number,
        )
